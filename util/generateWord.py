from docx import Document
# from docx.document import Document
from PyPDF2 import PdfReader

import os
from util.rutas import dir

class Word:
    pathDir:str
    dest = os.path.join(dir,"public")
    template = "Ficha de Inscripcion Josefina 2022-2023 (1).doc.pdf"
    def __init__(self,name:str) -> None:
        self.name = name
        self.document = Document()
        self.pathDir = self.path()




    def read_template(self):
        reader = PdfReader(f"{os.path.join(self.dest,self.template)}")
        return reader.pages[0].extract_text()

    def copy(self):
        self.document.add_paragraph(self.read_template())


    def write(self,**kwargs):
        self.document.add_paragraph(kwargs['p'])

    def path(self,path=os.path.join(dir,"public")):
        return f"{os.path.join(path,self.name)}"

    def save(self):
        self.document.save(self.pathDir)


